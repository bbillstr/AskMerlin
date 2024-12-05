from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for
import google.generativeai as genai
import os
from werkzeug.utils import secure_filename
import sqlite3
from datetime import datetime
import logging
import tiktoken  # For token counting
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from docx import Document
import pandas as pd
import io
import mimetypes
import time
from tenacity import retry, stop_after_attempt, wait_exponential
import openpyxl
from docx.shared import Inches, Pt

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['SECRET_KEY'] = os.urandom(24)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Configure Gemini AI
GOOGLE_API_KEY = ""  # Replace with your API key
logger.info("Configuring Gemini AI...")
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')
logger.info("Gemini AI configured successfully")

# Initialize tokenizer
tokenizer = tiktoken.get_encoding("cl100k_base")

def count_tokens(text: str) -> int:
    """Count the number of tokens in a text string."""
    return len(tokenizer.encode(text))

def log_token_usage(prompt: str, response: str):
    """Log token usage for both input and output."""
    input_tokens = count_tokens(prompt)
    output_tokens = count_tokens(response)
    total_tokens = input_tokens + output_tokens
    
    logger.info("Token Usage Summary:")
    logger.info(f"Input tokens:  {input_tokens:,}")
    logger.info(f"Output tokens: {output_tokens:,}")
    logger.info(f"Total tokens:  {total_tokens:,}")
    logger.info(f"Estimated cost: ${(total_tokens / 1000) * 0.0010:.4f} USD")

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def init_db():
    logger.info("Initializing database...")
    conn = sqlite3.connect('security_docs.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS documents
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
         title TEXT NOT NULL,
         content TEXT NOT NULL,
         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS chat_history
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
         question TEXT NOT NULL,
         answer TEXT NOT NULL,
         timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP)
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS feedback
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
         message_id INTEGER NOT NULL,
         feedback_type TEXT NOT NULL,
         timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP)
    ''')
    conn.commit()
    conn.close()
    logger.info("Database initialized successfully")

def get_document_content():
    logger.info("Retrieving all documents from database...")
    conn = sqlite3.connect('security_docs.db')
    c = conn.cursor()
    c.execute('SELECT content FROM documents')
    docs = c.fetchall()
    conn.close()
    combined_content = "\n".join([doc[0] for doc in docs])
    logger.info(f"Retrieved {len(docs)} documents. Total content length: {len(combined_content)} characters")
    return combined_content

@app.route('/')
def index():
    return redirect(url_for('chat'))

@app.route('/chat')
def chat():
    return render_template('chat.html')

@app.route('/documents')
def documents():
    return render_template('documents.html')

@app.route('/api/documents')
def list_documents():
    try:
        conn = sqlite3.connect('security_docs.db')
        c = conn.cursor()
        c.execute('''
            SELECT id, title, content, uploaded_at 
            FROM documents 
            ORDER BY uploaded_at DESC
        ''')
        docs = c.fetchall()
        conn.close()
        
        documents = [{
            'id': d[0],
            'title': d[1],
            'preview': d[2][:200] + '...' if len(d[2]) > 200 else d[2],
            'uploaded_at': d[3]
        } for d in docs]
        
        logger.info(f"Retrieved {len(documents)} documents")
        return jsonify(documents)
        
    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        return jsonify({'error': 'Failed to load documents'}), 500

@app.route('/documents/<int:doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    try:
        conn = sqlite3.connect('security_docs.db')
        c = conn.cursor()
        
        # Check if document exists
        c.execute('SELECT id FROM documents WHERE id = ?', (doc_id,))
        if not c.fetchone():
            conn.close()
            return jsonify({'error': 'Document not found'}), 404
            
        c.execute('DELETE FROM documents WHERE id = ?', (doc_id,))
        conn.commit()
        conn.close()
        
        logger.info(f"Document {doc_id} deleted successfully")
        return jsonify({'message': 'Document deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting document {doc_id}: {str(e)}")
        return jsonify({'error': 'Failed to delete document'}), 500

def get_relevant_context(question, all_content):
    # Use embeddings to find most relevant sections
    sections = all_content.split('\n\n')
    relevant_sections = []
    for section in sections:
        if semantic_similarity(question, section) > 0.7:
            relevant_sections.append(section)
    return '\n\n'.join(relevant_sections[:3])

def find_citations(answer, context):
    sentences = context.split('.')
    citations = []
    for sentence in sentences:
        if sentence.strip() in answer:
            citations.append(sentence.strip())
    return citations

def semantic_similarity(text1, text2):
    """Calculate semantic similarity between two texts using TF-IDF and cosine similarity."""
    vectorizer = TfidfVectorizer()
    try:
        tfidf_matrix = vectorizer.fit_transform([text1, text2])
        return cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    except:
        return 0

@app.route('/feedback', methods=['POST'])
def save_feedback():
    try:
        data = request.json
        message_id = data.get('messageId')
        feedback_type = data.get('type')
        
        conn = sqlite3.connect('security_docs.db')
        c = conn.cursor()
        c.execute('INSERT INTO feedback (message_id, feedback_type) VALUES (?, ?)',
                 (message_id, feedback_type))
        conn.commit()
        conn.close()
        
        return jsonify({'message': 'Feedback saved'})
    except Exception as e:
        logger.error(f"Error saving feedback: {str(e)}")
        return jsonify({'error': str(e)}), 500

def extract_questions_from_docx(file):
    """Extract questions and their paragraph locations from a Word document"""
    doc = Document(file)
    questions = []
    question_locations = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        # Check if the text ends with a question mark or contains question keywords
        if (text.endswith('?') or 
            text.lower().startswith(('what', 'how', 'why', 'when', 'where', 'which', 'can', 'could', 'would', 'will'))):
            if text:  # Only add non-empty questions
                questions.append(text)
                question_locations.append(i)
                logger.info(f"Found question at position {i}: {text[:50]}...")
    
    logger.info(f"Found {len(questions)} questions in Word document")
    return questions, question_locations

def update_word_document(file_content, questions_and_answers, question_locations):
    """Update the original Word document with answers after each question"""
    try:
        # Create a document from the original content
        doc_io = io.BytesIO(file_content)
        doc = Document(doc_io)
        
        # Create a dictionary mapping questions to answers
        qa_dict = {qa['question'].strip(): qa['answer'].strip() for qa in questions_and_answers}
        
        # Process questions in reverse order to maintain paragraph indices
        for i in range(len(question_locations) - 1, -1, -1):
            loc = question_locations[i]
            question = doc.paragraphs[loc].text.strip()
            
            if question in qa_dict:
                # Insert answer after the question
                # Add a new paragraph by inserting at the next index
                doc.paragraphs[loc].add_run('\n')  # Add newline after question
                
                # Create new paragraph for answer
                new_para = doc.add_paragraph()
                # Move the new paragraph to the correct position
                doc._body._body.insert(loc + 1, new_para._p)
                
                # Add styled "Answer: " prefix
                answer_run = new_para.add_run("Answer: ")
                answer_run.bold = True
                answer_run.font.size = Pt(11)
                answer_run.font.name = 'Calibri'
                
                # Add the answer text
                answer_text_run = new_para.add_run(qa_dict[question])
                answer_text_run.font.size = Pt(11)
                answer_text_run.font.name = 'Calibri'
                
                # Style the paragraph
                new_para.paragraph_format.left_indent = Inches(0.5)
                new_para.paragraph_format.space_before = Pt(6)
                new_para.paragraph_format.space_after = Pt(12)
                
                # Add a blank line after the answer
                blank_para = doc.add_paragraph()
                doc._body._body.insert(loc + 2, blank_para._p)
                
                logger.info(f"Added answer after question at position {loc}")
        
        # Save to buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        logger.info("Successfully updated Word document with answers")
        return doc_buffer
        
    except Exception as e:
        logger.error(f"Error updating Word document: {str(e)}")
        raise

@app.route('/batch-questions', methods=['POST'])
def process_batch_questions():
    if 'document' not in request.files:
        return jsonify({'error': 'No document provided'}), 400
        
    file = request.files['document']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
        
    try:
        # Extract questions based on file type
        file_ext = file.filename.rsplit('.', 1)[1].lower()
        
        # Create a copy of the file in memory for multiple reads
        file_content = file.read()
        file_in_memory = io.BytesIO(file_content)
        
        logger.info(f"Processing {file_ext} file: {file.filename}")
        
        # Get questions based on file type
        if file_ext == 'docx':
            questions, question_locations = extract_questions_from_docx(io.BytesIO(file_content))
            logger.info(f"Extracted {len(questions)} questions from Word document")
            if not questions:
                return jsonify({'error': 'No questions found in the document'}), 400
        elif file_ext in ['xlsx', 'xls']:
            question_col = request.form.get('questionColumn')
            answer_col = request.form.get('answerColumn')
            
            if not question_col or not answer_col:
                return jsonify({'error': 'Question and answer columns must be specified'}), 400
                
            df = pd.read_excel(io.BytesIO(file_content))
            questions = df[question_col].dropna().tolist()
            logger.info(f"Extracted {len(questions)} questions from Excel document")
            
            if not questions:
                return jsonify({'error': 'No questions found in the specified column'}), 400
            
        # Get context from uploaded documents
        context = get_document_content()
        
        # Create a single prompt with all questions
        questions_text = "\n".join([f"Question {i+1}: {q}" for i, q in enumerate(questions)])
        
        prompt = f"""
        Context: {context}

        Please answer the following questions based on the provided documentation. 
        For each question, provide a clear and concise answer.
        If the information is not available in the context, respond with "Information not found in the documentation."

        Format your response exactly as follows:
        Answer 1: [Your answer to question 1]
        Answer 2: [Your answer to question 2]
        etc.

        Questions:
        {questions_text}
        """

        logger.info("Making single API call for all questions...")
        # Make a single API call
        response = generate_answer(prompt)
        logger.info("Received response from API")

        # Parse the responses
        responses = []
        try:
            answer_parts = response.split('Answer ')
            for part in answer_parts[1:]:
                try:
                    if ':' in part:
                        answer_text = part.split(':', 1)[1].strip()
                        question_index = len(responses)
                        
                        if question_index < len(questions):
                            responses.append({
                                'question': questions[question_index],
                                'answer': answer_text
                            })
                except Exception as e:
                    logger.error(f"Error parsing individual answer: {str(e)}")
            
            # Add any missing answers
            while len(responses) < len(questions):
                responses.append({
                    'question': questions[len(responses)],
                    'answer': "No answer generated for this question"
                })
            
        except Exception as e:
            logger.error(f"Error parsing responses: {str(e)}")
            return jsonify({'error': 'Failed to parse responses'}), 500

        # Create or update document with answers
        try:
            if file_ext == 'docx':
                file_buffer = update_word_document(file_content, responses, question_locations)
            else:
                file_buffer = update_excel_with_answers(file_content, responses, question_col, answer_col)
            
            # Store the file buffer in a temporary file
            temp_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 
                                        f'response_{datetime.now().strftime("%Y%m%d_%H%M%S")}.{file_ext}')
            with open(temp_file_path, 'wb') as f:
                f.write(file_buffer.getvalue())
            
            return jsonify({
                'message': f'Processed {len(questions)} questions successfully',
                'results': responses,
                'downloadUrl': f'/download-responses/{os.path.basename(temp_file_path)}'
            })
                
        except Exception as e:
            logger.error(f"Error creating response document: {str(e)}")
            return jsonify({
                'message': f'Processed {len(questions)} questions successfully',
                'results': responses,
                'error': 'Could not create downloadable document'
            })
            
    except Exception as e:
        logger.error(f"Error processing questions: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

def update_excel_with_answers(file_content, questions_and_answers, question_col, answer_col):
    """Update Excel file with answers while preserving formatting"""
    try:
        # Load workbook with formatting
        workbook = openpyxl.load_workbook(io.BytesIO(file_content))
        worksheet = workbook.active
        
        # Read data into DataFrame for easier question matching
        df = pd.read_excel(io.BytesIO(file_content))
        
        # Create a dictionary mapping questions to answers
        qa_dict = {str(qa['question']).strip(): qa['answer'] for qa in questions_and_answers}
        
        # Get column indices (1-based for openpyxl)
        answer_col_idx = df.columns.get_loc(answer_col) + 1
        question_col_idx = df.columns.get_loc(question_col) + 1
        
        # Update only the answer cells where questions match
        for row_idx, row in enumerate(worksheet.iter_rows(min_row=2), start=2):  # Start from row 2 to skip header
            question = str(row[question_col_idx - 1].value).strip()
            if question in qa_dict:
                # Only update the answer cell, preserving all other formatting
                answer_cell = worksheet.cell(row=row_idx, column=answer_col_idx)
                answer_cell.value = qa_dict[question]
                logger.info(f"Updated answer for question: {question[:50]}...")
        
        # Save to buffer while preserving all formatting
        buffer = io.BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        
        logger.info("Successfully updated Excel document with answers while preserving formatting")
        return buffer
        
    except Exception as e:
        logger.error(f"Error updating Excel document: {str(e)}")
        raise

@app.route('/download-responses/<filename>')
def download_responses(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
            
        if filename.endswith('docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            
        response = send_file(
            file_path,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
        
        # Clean up the temporary file after sending
        @response.call_on_close
        def cleanup():
            try:
                os.remove(file_path)
            except Exception as e:
                logger.error(f"Error cleaning up temporary file: {str(e)}")
                
        return response
        
    except Exception as e:
        logger.error(f"Error downloading response file: {str(e)}")
        return jsonify({'error': 'Failed to download response file'}), 500

@app.route('/get-columns', methods=['POST'])
def get_columns():
    if 'document' not in request.files:
        return jsonify({'error': 'No document provided'}), 400
        
    file = request.files['document']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
        
    try:
        df = pd.read_excel(file)
        columns = df.columns.tolist()
        return jsonify({'columns': columns})
    except Exception as e:
        logger.error(f"Error reading Excel columns: {str(e)}")
        return jsonify({'error': 'Failed to read Excel columns'}), 500

@app.route('/upload', methods=['POST'])
def upload_document():
    logger.info("=== New Document Upload ===")
    
    if 'document' not in request.files and 'content' not in request.form:
        logger.error("No document provided in request")
        return jsonify({'error': 'No document provided'}), 400

    try:
        content = ""
        if 'document' in request.files:
            file = request.files['document']
            if file.filename:
                if not allowed_file(file.filename):
                    return jsonify({'error': 'File type not supported'}), 400

                logger.info(f"Processing file: {file.filename}")
                filename = secure_filename(file.filename)
                file_ext = filename.rsplit('.', 1)[1].lower()

                # Process different file types
                try:
                    if file_ext in ['docx']:
                        content = extract_text_from_docx(file)
                    elif file_ext in ['xlsx', 'xls']:
                        content = extract_text_from_excel(file)
                    else:  # txt files
                        content = file.read().decode('utf-8')
                except Exception as e:
                    logger.error(f"Error processing file: {str(e)}")
                    return jsonify({'error': f'Error processing file: {str(e)}'}), 400

                # Log token count for uploaded file
                token_count = count_tokens(content)
                logger.info(f"File processed. Size: {len(content):,} chars, {token_count:,} tokens")
        else:
            content = request.form['content']
            token_count = count_tokens(content)
            logger.info(f"Text input received. Size: {len(content):,} chars, {token_count:,} tokens")

        # Check if content is too large
        if token_count > 100000:  # Gemini's approximate token limit
            logger.warning(f"Content too large: {token_count:,} tokens (limit: 100,000)")
            return jsonify({'error': 'Content too large. Please reduce the size.'}), 400

        conn = sqlite3.connect('security_docs.db')
        c = conn.cursor()
        title = request.form.get('title', 'Untitled')
        c.execute('INSERT INTO documents (title, content) VALUES (?, ?)',
                 (title, content))
        conn.commit()
        conn.close()
        logger.info(f"Document '{title}' saved successfully")

        return jsonify({'message': 'Document uploaded successfully'})
    except Exception as e:
        logger.error(f"Upload error: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=4, max=10),
    retry_error_callback=lambda _: None
)
def generate_answer(prompt):
    """Generate answer with retry logic"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        logger.error(f"Error generating response: {str(e)}")
        return "Error: Unable to generate response due to API limitations. Please try again later."

@app.route('/ask', methods=['POST'])
def ask_question():
    logger.info("=== New Question Request ===")
    try:
        question = request.json.get('question')
        if not question:
            logger.error("No question provided in request")
            return jsonify({'error': 'No question provided'}), 400

        logger.info(f"Question: {question}")
        
        # Get all document content
        context = get_document_content()
        logger.info(f"Retrieved context: {len(context)} characters")
        
        # Construct prompt
        prompt = f"""
        Context: {context}
        
        Question: {question}
        
        Please provide a clear and concise answer based on the provided documentation. 
        If relevant, include:
        - Specific requirements or policies
        - Best practices and industry standards
        - Related security considerations
        
        If the information is not available in the context, respond with "I cannot find relevant information in the provided documentation."
        """
        
        # Log pre-request token count
        logger.info("Sending request to Gemini AI...")
        logger.info(f"Input prompt tokens: {count_tokens(prompt):,}")
        
        # Generate response using Gemini
        response = generate_answer(prompt)
        
        # Log token usage
        log_token_usage(prompt, response)
        
        logger.info("Response Preview:")
        logger.info(f"{response[:200]}...")  # Log first 200 chars
        
        return jsonify({'answer': response})
    except Exception as e:
        logger.error(f"Error processing question: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    init_db()
    logger.info("Starting Flask application...")
    app.run(debug=True)
